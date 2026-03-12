import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from io import BytesIO
import io

def sugerir_descripcion(variable, tipo_dato):
    nombre = str(variable).lower()

    if "id" in nombre:
        return "Identificador único del registro o de la entidad asociada."
    elif "nombre" in nombre:
        return "Nombre asociado al registro."
    elif "apellido" in nombre:
        return "Apellido asociado al registro."
    elif "edad" in nombre:
        return "Edad de la persona registrada en la base de datos."
    elif "sexo" in nombre or "genero" in nombre or "género" in nombre:
        return "Sexo o género reportado en el registro."
    elif "fecha" in nombre:
        return "Fecha asociada al evento o registro."
    elif "hora" in nombre:
        return "Hora asociada al evento o registro."
    elif "correo" in nombre or "email" in nombre:
        return "Correo electrónico asociado al registro."
    elif "telefono" in nombre or "teléfono" in nombre or "celular" in nombre:
        return "Número telefónico asociado al registro."
    elif "direccion" in nombre or "dirección" in nombre:
        return "Dirección asociada al registro."
    elif "comuna" in nombre:
        return "Comuna reportada en la base de datos."
    elif "barrio" in nombre:
        return "Barrio reportado en la base de datos."
    elif "municipio" in nombre:
        return "Municipio asociado al registro."
    elif "departamento" in nombre:
        return "Departamento asociado al registro."
    elif "zona" in nombre:
        return "Zona geográfica o administrativa asociada al registro."
    elif "punto" in nombre:
        return "Punto de atención, sede o ubicación asociada al registro."
    elif "servicio" in nombre:
        return "Tipo de servicio asociado al registro."
    elif "estado" in nombre:
        return "Estado o situación reportada en el registro."
    elif "etnia" in nombre:
        return "Pertenencia étnica reportada en el registro."
    elif "condicion" in nombre or "condición" in nombre:
        return "Condición o característica reportada para el registro."
    elif "dispositivo" in nombre:
        return "Identificador o referencia del dispositivo asociado al registro."
    elif "trafico" in nombre or "tráfico" in nombre:
        return "Volumen de tráfico de datos registrado."
    elif "sesion" in nombre or "sesión" in nombre:
        return "Información relacionada con la sesión registrada."
    elif "usuario" in nombre:
        return "Cantidad o referencia de usuarios asociados al registro."
    else:
        if tipo_dato == "Numérico":
            return "Variable numérica registrada en la base de datos."
        elif tipo_dato == "Fecha":
            return "Variable de fecha registrada en la base de datos."
        else:
            return "Variable de texto registrada en la base de datos."
def clasificar_variable(variable, tipo_dato):
    nombre = str(variable).lower()

    # Identificadoras
    if any(p in nombre for p in [
        "id", "identificador", "codigo", "código", "documento",
        "cedula", "cédula", "nombre", "apellido", "correo",
        "email", "telefono", "teléfono", "celular", "direccion", "dirección"
    ]):
        return "Identificadora"

    # Temporales
    elif any(p in nombre for p in [
        "fecha", "hora", "periodo", "período", "vigencia", "anio", "año", "mes", "dia", "día"
    ]):
        return "Temporal"

    # Geográficas
    elif any(p in nombre for p in [
        "pais", "país", "departamento", "municipio", "ciudad", "comuna",
        "barrio", "zona", "direccion", "dirección", "latitud", "longitud", "sede", "punto"
    ]):
        return "Geográfica"

    # Sensibles
    elif any(p in nombre for p in [
        "salud", "discapacidad", "etnia", "sexo", "genero", "género",
        "victima", "víctima", "migrante", "desplazado", "condicion", "condición"
    ]):
        return "Sensible"

    # Cuantitativas
    elif tipo_dato == "Numérico":
        return "Cuantitativa"

    # Temporales por tipo técnico
    elif tipo_dato == "Fecha":
        return "Temporal"

    # Por defecto
    else:
        return "Cualitativa"
        
def convertir_diccionario_a_excel(df_diccionario):
    output = BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_diccionario.to_excel(writer, index=False, sheet_name="Diccionario")

    output.seek(0)
    return output

col1, col2, col3 = st.columns([1,2,1])

with col2:
    st.image("SIRDATA.png", width=1800)

st.subheader("Sistema de Revisión de Datos")
st.write("Observatorio de Bases de Datos - Sistema Estadístico Municipal de Pereira")
st.markdown("""
### Finalidad y alcance

Esta plataforma permite analizar la estructura de las bases de datos, generar diccionarios de datos, evaluar indicadores de calidad de la información y analizar condiciones de anonimización, facilitando el aprovechamiento estadístico de los registros administrativos.
""")


st.header("Cargar base de datos")

archivo = st.file_uploader("Sube un archivo Excel o CSV", type=["xlsx", "csv", "json"])



if archivo is not None:

    if archivo.name.endswith(".csv"):
        df = pd.read_csv(archivo)

    elif archivo.name.endswith(".xlsx"):
        df = pd.read_excel(archivo, engine="openpyxl")

    elif archivo.name.endswith(".json"):
        df = pd.read_json(archivo)

    st.success("Archivo cargado correctamente")

    st.write("Número de registros:", df.shape[0])
    st.write("Número de variables:", df.shape[1])

    st.subheader("Vista previa de la base")
    st.dataframe(df.head())

    st.header("Perfilamiento de variables")

perfil = []

for col in df.columns:
    tipo = str(df[col].dtype)
    nulos = df[col].isnull().sum()
    porcentaje_nulos = round((nulos / len(df)) * 100, 2)

    serie_sin_nulos = df[col].dropna()

    # Convertir valores complejos (dict, list) a texto para poder analizarlos
    serie_para_unicos = serie_sin_nulos.apply(
        lambda x: str(x) if isinstance(x, (dict, list)) else x
    )

    unicos = serie_para_unicos.nunique()

    ejemplo = ""
    if not serie_sin_nulos.empty:
        ejemplo = serie_sin_nulos.iloc[0]
        if isinstance(ejemplo, (dict, list)):
            ejemplo = str(ejemplo)

    perfil.append({
        "Variable": col,
        "Tipo de dato": tipo,
        "Valores únicos": unicos,
        "Nulos": nulos,
        "% Nulos": porcentaje_nulos,
        "Ejemplo": ejemplo
    })

perfil_df = pd.DataFrame(perfil)

st.dataframe(perfil_df)

st.header("Generación automática de diccionario de datos")

diccionario = []

import json

def normalizar_valor(valor):
    if isinstance(valor, (dict, list)):
        return json.dumps(valor, ensure_ascii=False, sort_keys=True)
    return valor

for col in df.columns:
    tipo = df[col].dtype
    
    serie = df[col].apply(normalizar_valor)
    valores_unicos = serie.nunique(dropna=True)

    ejemplo = df[col].dropna().astype(str).head(1)

    if len(ejemplo) > 0:
        ejemplo = ejemplo.values[0]
    else:
        ejemplo = ""

    if "int" in str(tipo) or "float" in str(tipo):
        tipo_dato = "Numérico"
    elif "datetime" in str(tipo):
        tipo_dato = "Fecha"
    else:
        tipo_dato = "Texto"

    diccionario.append({
        "Variable": col,
        "Clasificación sugerida": clasificar_variable(col, tipo_dato),
        "Tipo de dato": tipo_dato,
        "Tipo técnico": tipo,
        "Valores nulos": nulos,
        "% nulos": porcentaje_nulos,
        "Valores únicos": valores_unicos,
        "Ejemplo": ejemplo,
        "Descripción sugerida": sugerir_descripcion(col, tipo_dato)
    })

diccionario_df = pd.DataFrame(diccionario)

st.dataframe(diccionario_df)

excel_diccionario = convertir_diccionario_a_excel(diccionario_df)

st.download_button(
    label="Descargar diccionario de datos en Excel",
    data=excel_diccionario,
    file_name="diccionario_datos_sirdata.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

st.header("Indicadores de calidad de datos")

total_registros = df.shape[0]
total_variables = df.shape[1]

# -------------------------
# 1. COMPLETITUD
# -------------------------
campos_esperados = total_registros * total_variables
campos_diligenciados = df.notnull().sum().sum()
completitud = round((campos_diligenciados / campos_esperados) * 100, 2)

# -------------------------
# 2. UNICIDAD
# -------------------------

df_duplicados = df.copy()

for col in df_duplicados.columns:
    df_duplicados[col] = df_duplicados[col].apply(
        lambda x: str(x) if isinstance(x, (dict, list)) else x
    )

duplicados = df_duplicados.duplicated().sum()
registros_unicos = total_registros - duplicados
unicidad = round((registros_unicos / total_registros) * 100, 2)

# -------------------------
# 3. EXACTITUD (aproximación inicial)
# -------------------------
valores_validos = df.notnull().sum().sum()
exactitud = round((valores_validos / campos_esperados) * 100, 2)

# -------------------------
# 4. CONSISTENCIA (aproximación inicial)
# -------------------------
# -------------------------
# 4. CONSISTENCIA (mejorada)
# -------------------------
columnas_consistentes = 0

for col in df.columns:
    serie = df[col].dropna()
    
    if len(serie) == 0:
        columnas_consistentes += 1
    else:
        tipos_detectados = serie.map(type).nunique()
        if tipos_detectados == 1:
            columnas_consistentes += 1

consistencia = round((columnas_consistentes / total_variables) * 100, 2)

# -------------------------
# 5. VALIDEZ (aproximación inicial)
# -------------------------
valores_validos_tipo = 0

for col in df.columns:
    serie = df[col].dropna()
    valores_validos_tipo += len(serie)

validez = round((valores_validos_tipo / campos_esperados) * 100, 2)

# -------------------------
# Función para asignar color
# -------------------------
def clasificar_color(valor):
    if valor >= 80:
        return "Alta"
    elif valor >= 65:
        return "Media"
    else:
        return "Baja"

# -------------------------
# Mostrar métricas
# -------------------------
col1, col2, col3, col4, col5 = st.columns(5)

col1.metric("Completitud", f"{completitud}%")
col2.metric("Exactitud", f"{exactitud}%")
col3.metric("Consistencia", f"{consistencia}%")
col4.metric("Unicidad", f"{unicidad}%")
col5.metric("Validez", f"{validez}%")

# -------------------------
# Tabla de indicadores
# -------------------------
indicadores_df = pd.DataFrame({
    "Indicador": ["Completitud", "Exactitud", "Consistencia", "Unicidad", "Validez"],
    "Valor": [completitud, exactitud, consistencia, unicidad, validez]
})

indicadores_df["Nivel"] = indicadores_df["Valor"].apply(clasificar_color)

# -------------------------
# Gráfico de barras con colores
# -------------------------
fig = px.bar(
    indicadores_df,
    x="Indicador",
    y="Valor",
    color="Nivel",
    text="Valor",
    color_discrete_map={
        "Alta": "green",
        "Media": "gold",
        "Baja": "red"
    },
    title="Indicadores de calidad de datos"
)

fig.update_traces(texttemplate='%{text:.2f}%', textposition='outside')
fig.update_layout(
    yaxis_title="Porcentaje",
    xaxis_title="Indicador",
    yaxis_range=[0, 110]
)

st.plotly_chart(fig, use_container_width=True)

st.markdown("### Plan de mejora de calidad de datos")

plan_mejora = []

def clasificar_nivel(valor):
    if valor >= 80:
        return "Alta"
    elif valor >= 65:
        return "Media"
    else:
        return "Baja"

# Completitud
nivel_completitud = clasificar_nivel(completitud)
if completitud < 100:
    accion_completitud = "Revisar campos vacíos, completar información faltante y definir cuáles variables son obligatorias y cuáles opcionales."
else:
    accion_completitud = "No se requieren acciones. El indicador ya alcanzó el 100%."

plan_mejora.append({
    "Indicador": "Completitud",
    "Resultado": f"{completitud}%",
    "Nivel": nivel_completitud,
    "Acción sugerida": accion_completitud
})

# Exactitud
nivel_exactitud = clasificar_nivel(exactitud)
if exactitud < 100:
    accion_exactitud = "Revisar valores erróneos, estandarizar reglas de validación y depurar errores de digitación o captura."
else:
    accion_exactitud = "No se requieren acciones. El indicador ya alcanzó el 100%."

plan_mejora.append({
    "Indicador": "Exactitud",
    "Resultado": f"{exactitud}%",
    "Nivel": nivel_exactitud,
    "Acción sugerida": accion_exactitud
})

# Consistencia
nivel_consistencia = clasificar_nivel(consistencia)
if consistencia < 100:
    accion_consistencia = "Revisar coherencia entre variables, definir reglas de negocio y validar que la estructura de los registros sea homogénea."
else:
    accion_consistencia = "No se requieren acciones. El indicador ya alcanzó el 100%."

plan_mejora.append({
    "Indicador": "Consistencia",
    "Resultado": f"{consistencia}%",
    "Nivel": nivel_consistencia,
    "Acción sugerida": accion_consistencia
})

# Unicidad
nivel_unicidad = clasificar_nivel(unicidad)
if unicidad < 100:
    accion_unicidad = "Identificar y eliminar registros duplicados, y definir una llave única para controlar futuros cargues."
else:
    accion_unicidad = "No se requieren acciones. El indicador ya alcanzó el 100%."

plan_mejora.append({
    "Indicador": "Unicidad",
    "Resultado": f"{unicidad}%",
    "Nivel": nivel_unicidad,
    "Acción sugerida": accion_unicidad
})

# Validez
nivel_validez = clasificar_nivel(validez)
if validez < 100:
    accion_validez = "Estandarizar formatos, revisar tipos de dato y validar los valores frente a catálogos o dominios permitidos."
else:
    accion_validez = "No se requieren acciones. El indicador ya alcanzó el 100%."

plan_mejora.append({
    "Indicador": "Validez",
    "Resultado": f"{validez}%",
    "Nivel": nivel_validez,
    "Acción sugerida": accion_validez
})

plan_mejora_df = pd.DataFrame(plan_mejora)

st.markdown("### Diagnóstico general de calidad")

indicadores_bajos = []
indicadores_medios = []

for fila in plan_mejora:
    if fila["Nivel"] == "Baja":
        indicadores_bajos.append(fila["Indicador"])
    elif fila["Nivel"] == "Media":
        indicadores_medios.append(fila["Indicador"])

if len(indicadores_bajos) == 0 and len(indicadores_medios) == 0:
    st.success("La base presenta una calidad alta en todos los indicadores evaluados.")
elif len(indicadores_bajos) == 0:
    st.warning(
        f"La base presenta calidad aceptable, pero se recomienda fortalecer los indicadores en nivel medio: {', '.join(indicadores_medios)}."
    )
else:
    st.error(
        f"La base presenta oportunidades de mejora importantes. Se recomienda priorizar acciones sobre los siguientes indicadores: {', '.join(indicadores_bajos)}."
    )
st.dataframe(plan_mejora_df)

st.header("Evaluación de anonimización")

# Listas de palabras clave
identificadores_directos = [
    "nombre", "apellido", "cedula", "cédula", "documento",
    "telefono", "teléfono", "celular", "correo", "email",
    "direccion", "dirección"
]

identificadores_indirectos = [
    "edad", "fecha_nacimiento", "nacimiento", "barrio",
    "comuna", "municipio", "institucion", "institución",
    "sede", "punto", "zona"
]

variables_sensibles = [
    "salud", "discapacidad", "victima", "víctima", "etnia",
    "genero", "género", "sexo", "migrante", "desplazado",
    "desplazados", "condicion", "condición"
]

# Clasificación de variables
directas_detectadas = []
indirectas_detectadas = []
sensibles_detectadas = []

for col in df.columns:
    for col in df.columns:
        col_lower = str(col).lower()

    if any(palabra in col_lower for palabra in identificadores_directos):
        directas_detectadas.append(col)

    if any(palabra in col_lower for palabra in identificadores_indirectos):
        indirectas_detectadas.append(col)

    if any(palabra in col_lower for palabra in variables_sensibles):
        sensibles_detectadas.append(col)

# Mostrar resultados
col1, col2, col3 = st.columns(3)

with col1:
    st.subheader("Identificadores directos")
    if directas_detectadas:
        for v in directas_detectadas:
            st.write(f"- {v}")
    else:
        st.write("No se detectaron")

with col2:
    st.subheader("Identificadores indirectos")
    if indirectas_detectadas:
        for v in indirectas_detectadas:
            st.write(f"- {v}")
    else:
        st.write("No se detectaron")

with col3:
    st.subheader("Variables sensibles")
    if sensibles_detectadas:
        for v in sensibles_detectadas:
            st.write(f"- {v}")
    else:
        st.write("No se detectaron")

# Recomendación de publicación
st.subheader("Recomendación de publicación")

hallazgos = []
acciones = []

if len(directas_detectadas) > 0:
    recomendacion = "Acceso restringido"
    mensaje = "La base no se recomienda para publicación en su estado actual."

    for var in directas_detectadas:
        hallazgos.append(f"La variable '{var}' fue identificada como identificador directo.")
        acciones.append(f"Eliminar o anonimizar la variable '{var}' antes de publicar la base.")

    motivo = "La presencia de identificadores directos permite identificar personas de manera individual."

elif len(sensibles_detectadas) > 0 or len(indirectas_detectadas) > 0:
    recomendacion = "Publicable con anonimización previa"
    mensaje = "La base requiere ajustes antes de su publicación."

    for var in indirectas_detectadas:
        hallazgos.append(f"La variable '{var}' fue identificada como identificador indirecto.")
        acciones.append(f"Revisar la variable '{var}' y evaluar su generalización o agregación antes de publicar.")

    for var in sensibles_detectadas:
        hallazgos.append(f"La variable '{var}' fue identificada como variable sensible.")
        acciones.append(f"Evaluar si la variable '{var}' debe eliminarse, agruparse o restringirse antes de publicar.")

    motivo = "La combinación de variables indirectas o sensibles puede generar riesgo de reidentificación o exposición de información protegida."

else:
    recomendacion = "Publicable"
    mensaje = "La base no presenta identificadores directos ni variables sensibles evidentes."
    motivo = "No se detectaron restricciones evidentes para su publicación desde la revisión automática inicial."

st.write(f"**Clasificación:** {recomendacion}")
st.info(mensaje)

st.markdown("### Motivo técnico")
st.write(motivo)

if hallazgos:
    st.markdown("### Hallazgos detectados")
    for h in hallazgos:
        st.write(f"- {h}")

if acciones:
    st.markdown("### Recomendaciones de corrección")
    acciones_unicas = list(dict.fromkeys(acciones))
    for a in acciones_unicas:
        st.write(f"- {a}")

        st.markdown("### Plan de anonimización sugerido")

plan = []

# Identificadores directos
for var in directas_detectadas:
    plan.append({
        "Variable": var,
        "Tipo de riesgo": "Identificador directo",
        "Acción sugerida": "Eliminar la variable o reemplazar por un identificador anónimo"
    })

# Identificadores indirectos
for var in indirectas_detectadas:
    plan.append({
        "Variable": var,
        "Tipo de riesgo": "Identificador indirecto",
        "Acción sugerida": "Generalizar o agrupar la información (ejemplo: edad por rangos, ubicación por zona)"
    })

# Variables sensibles
for var in sensibles_detectadas:
    plan.append({
        "Variable": var,
        "Tipo de riesgo": "Variable sensible",
        "Acción sugerida": "Evaluar eliminación, agregación estadística o restricción de acceso"
    })

if plan:
    plan_df = pd.DataFrame(plan)
    st.dataframe(plan_df)
else:
    st.success("No se requieren acciones de anonimización para esta base.")

def generar_reporte_word(
    nombre_archivo,
    total_registros,
    total_variables,
    completitud,
    exactitud,
    consistencia,
    unicidad,
    validez,
    plan_mejora_df,
    directas_detectadas,
    indirectas_detectadas,
    sensibles_detectadas,
    recomendacion,
    motivo,
    plan_df,
    diccionario_df
):
    doc = Document()

    doc.add_heading("Informe técnico de revisión de bases de datos", 0)
    doc.add_paragraph("SIRDATA – Sistema de Revisión de Datos")
    doc.add_paragraph("Observatorio de Bases de Datos – Sistema Estadístico Municipal de Pereira")

    # Información general
    doc.add_heading("1. Información general de la base", level=1)
    doc.add_paragraph(f"Nombre del archivo: {nombre_archivo}")
    doc.add_paragraph(f"Número de registros: {total_registros}")
    doc.add_paragraph(f"Número de variables: {total_variables}")

    # Indicadores
    doc.add_heading("2. Indicadores de calidad de datos", level=1)
    doc.add_paragraph(f"Completitud: {completitud}%")
    doc.add_paragraph(f"Exactitud: {exactitud}%")
    doc.add_paragraph(f"Consistencia: {consistencia}%")
    doc.add_paragraph(f"Unicidad: {unicidad}%")
    doc.add_paragraph(f"Validez: {validez}%")

    # Plan de mejora
    doc.add_heading("3. Plan de mejora de calidad de datos", level=1)
    tabla_mejora = doc.add_table(rows=1, cols=len(plan_mejora_df.columns))
    tabla_mejora.style = "Table Grid"
    hdr = tabla_mejora.rows[0].cells
    for i, col in enumerate(plan_mejora_df.columns):
        hdr[i].text = str(col)

    for _, row in plan_mejora_df.iterrows():
        celdas = tabla_mejora.add_row().cells
        for i, valor in enumerate(row):
            celdas[i].text = str(valor)

    # Anonimización
    doc.add_heading("4. Evaluación de anonimización", level=1)
    doc.add_paragraph(f"Clasificación: {recomendacion}")
    doc.add_paragraph(f"Motivo técnico: {motivo}")

    doc.add_paragraph("Identificadores directos detectados:")
    if directas_detectadas:
        for v in directas_detectadas:
            doc.add_paragraph(f"- {v}")
    else:
        doc.add_paragraph("- No se detectaron")

    doc.add_paragraph("Identificadores indirectos detectados:")
    if indirectas_detectadas:
        for v in indirectas_detectadas:
            doc.add_paragraph(f"- {v}")
    else:
        doc.add_paragraph("- No se detectaron")

    doc.add_paragraph("Variables sensibles detectadas:")
    if sensibles_detectadas:
        for v in sensibles_detectadas:
            doc.add_paragraph(f"- {v}")
    else:
        doc.add_paragraph("- No se detectaron")

    # Plan de anonimización
    doc.add_heading("5. Plan de anonimización sugerido", level=1)
    if not plan_df.empty:
        tabla_plan = doc.add_table(rows=1, cols=len(plan_df.columns))
        tabla_plan.style = "Table Grid"
        hdr = tabla_plan.rows[0].cells
        for i, col in enumerate(plan_df.columns):
            hdr[i].text = str(col)

        for _, row in plan_df.iterrows():
            celdas = tabla_plan.add_row().cells
            for i, valor in enumerate(row):
                celdas[i].text = str(valor)
    else:
        doc.add_paragraph("No se requieren acciones de anonimización para esta base.")

    # Diccionario
    doc.add_heading("6. Diccionario de datos", level=1)
    tabla_dic = doc.add_table(rows=1, cols=len(diccionario_df.columns))
    tabla_dic.style = "Table Grid"
    hdr = tabla_dic.rows[0].cells
    for i, col in enumerate(diccionario_df.columns):
        hdr[i].text = str(col)

    for _, row in diccionario_df.iterrows():
        celdas = tabla_dic.add_row().cells
        for i, valor in enumerate(row):
            celdas[i].text = str(valor)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

st.header("Reporte técnico final")

reporte_word = generar_reporte_word(
    nombre_archivo=archivo.name,
    total_registros=total_registros,
    total_variables=total_variables,
    completitud=completitud,
    exactitud=exactitud,
    consistencia=consistencia,
    unicidad=unicidad,
    validez=validez,
    plan_mejora_df=plan_mejora_df,
    directas_detectadas=directas_detectadas,
    indirectas_detectadas=indirectas_detectadas,
    sensibles_detectadas=sensibles_detectadas,
    recomendacion=recomendacion,
    motivo=motivo,
    plan_df=plan_df if 'plan_df' in locals() else pd.DataFrame(),
    diccionario_df=diccionario_df
)

st.download_button(
    label="Descargar informe técnico en Word",
    data=reporte_word,
    file_name="informe_tecnico_sirdata.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)